# import streamlit as st
# import requests
#
# # --- Konfiguration ---
# API_URL = "http://127.0.0.1:8000/analyze" # Sicherstellen, dass Port und Pfad zum Backend passen
#
# # --- Seitenaufbau ---
# st.set_page_config(page_title="AI CV Evaluator", layout="centered", initial_sidebar_state="expanded")
# st.title("📄 AI Lebenslauf-Bewerter")
# st.markdown("Lade deinen Lebenslauf hoch und erhalte eine KI-basierte Bewertung, wie gut er zu deiner Wunschrolle passt.")
#
# # --- Sidebar mit Filtern ---
# with st.sidebar:
#     st.header("🔎 Filteroptionen")
#     selected_role = st.selectbox(
#         "Jobrolle",
#         ["Alle", "Data Analyst", "Softwareentwickler", "IT-Consultant", "Projektmanager", "Marketing Specialist"]
#     )
#     selected_industry = st.selectbox(
#         "Branche",
#         ["Alle", "IT", "Marketing", "Logistik", "Beratung", "Finanzen"]
#     )
#     language = st.selectbox("Analyse-Sprache", ["de", "en", "fr"])
#
# # --- Kernfunktionen mit Caching für bessere Performance ---
# @st.cache_data
# def read_file(file):
#     """Liest den Inhalt verschiedener Dateitypen. Wird dank Cache nur einmal pro Datei ausgeführt."""
#     st.info(f"Lese Datei: {file.name}...")
#     if file.name.endswith(".txt"):
#         return file.read().decode("utf-8")
#     elif file.name.endswith(".pdf"):
#         # Hinweis: PyPDF2 muss installiert sein (pip install PyPDF2)
#         import PyPDF2 # Import hier, um die Startzeit der App nicht zu verlangsamen
#         reader = PyPDF2.PdfReader(file)
#         return "".join([page.extract_text() or "" for page in reader.pages])
#     elif file.name.endswith(".docx"):
#         # Hinweis: python-docx muss installiert sein (pip install python-docx)
#         import docx # Import hier, um die Startzeit der App nicht zu verlangsamen
#         doc = docx.Document(file)
#         return "\n".join([para.text for para in doc.paragraphs])
#     else:
#         st.error("Nicht unterstütztes Dateiformat.")
#         return ""
#
# @st.cache_data
# def analyze_cv(_text_content, language, role=None, industry=None):
#     """
#     Sendet den Text an die Backend-API. Wird dank Cache nur einmal für die gleiche Kombination
#     aus Text und Filtern ausgeführt.
#     """
#     payload = {
#         "text": _text_content,
#         "language": language,
#         "role": role,
#         "industry": industry
#     }
#     try:
#         # Setzt ein Timeout, um zu verhindern, dass die App bei Problemen einfriert
#         response = requests.post(API_URL, json=payload, timeout=90)
#         response.raise_for_status() # Löst einen Fehler für HTTP-Statuscodes wie 4xx oder 5xx aus
#         return True, response.json().get("result", "Keine gültige Antwort vom Server erhalten.")
#     except requests.exceptions.RequestException as e:
#         # Fängt Netzwerkfehler und fehlerhafte HTTP-Statuscodes ab
#         error_message = f"Fehler bei der API-Anfrage: {e}"
#         try:
#             # Versuche, die spezifische Fehlermeldung vom Server zu bekommen
#             error_detail = e.response.json().get("detail", "Keine Details.")
#             error_message = f"Fehler vom Server: {error_detail}"
#         except:
#             pass # Wenn die Antwort kein JSON ist, bleibe bei der allgemeinen Fehlermeldung
#         return False, error_message
#
# # --- Hauptlogik der App ---
# uploaded_file = st.file_uploader("📤 Lebenslauf hochladen (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"])
#
# if uploaded_file:
#     content = read_file(uploaded_file) # Nutzt die gecachte Funktion
#     if not content.strip():
#         st.warning("⚠️ Die Datei scheint leer zu sein oder konnte nicht gelesen werden.")
#     else:
#         st.text_area("📄 Vorschau des Lebenslaufs:", content, height=200)
#         if st.button("🧠 Jetzt analysieren"):
#             with st.spinner("Analyse läuft…"):
#                 # Nutzt die gecachte Funktion
#                 success, result = analyze_cv(
#                     content,
#                     language,
#                     None if selected_role == "Alle" else selected_role,
#                     None if selected_industry == "Alle" else selected_industry
#                 )
#
#             if success:
#                 st.success("✅ Analyse abgeschlossen!")
#                 st.markdown(f"### Ergebnis:\n> {result}")
#             else:
#                 st.error(result) # Zeigt die formatierte Fehlermeldung an
# else:
#     st.info("🔔 Bitte lade einen Lebenslauf hoch, um die Analyse zu starten.")

import streamlit as st

# --- Seitenaufbau ---
st.set_page_config(page_title="AI CV Evaluator (Dummy)", layout="centered", initial_sidebar_state="expanded")
st.title("📄 AI Lebenslauf-Bewerter (Demo-Modus)")
st.markdown("Lade deinen Lebenslauf hoch und erhalte eine simulierte KI-Bewertung.")

# --- Sidebar mit Filtern ---
with st.sidebar:
    st.header("🔎 Filteroptionen")
    selected_role = st.selectbox(
        "Jobrolle",
        ["Alle", "Data Analyst", "Softwareentwickler", "IT-Consultant", "Projektmanager", "Marketing Specialist"]
    )
    selected_industry = st.selectbox(
        "Branche",
        ["Alle", "IT", "Marketing", "Logistik", "Beratung", "Finanzen"]
    )
    language = st.selectbox("Analyse-Sprache", ["de", "en", "fr"])

# --- Datei-Lese-Funktion ---
@st.cache_data
def read_file(file):
    st.info(f"Lese Datei: {file.name}...")
    if file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    elif file.name.endswith(".pdf"):
        import PyPDF2
        reader = PyPDF2.PdfReader(file)
        return "".join([page.extract_text() or "" for page in reader.pages])
    elif file.name.endswith(".docx"):
        import docx
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        st.error("Nicht unterstütztes Dateiformat.")
        return ""

# --- Dummy-Analyse-Funktion ---
def dummy_analyze_cv(_text_content, language, role=None, industry=None):
    st.info("✅ Dummy-Analyse durchgeführt.")
    return "Ergebnis: JA, Bewertung: 4"

# --- Hauptlogik ---
uploaded_file = st.file_uploader("📤 Lebenslauf hochladen (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"])

if uploaded_file:
    content = read_file(uploaded_file)
    if not content.strip():
        st.warning("⚠️ Die Datei scheint leer zu sein oder konnte nicht gelesen werden.")
    else:
        st.text_area("📄 Vorschau des Lebenslaufs:", content, height=200)
        if st.button("🧠 Jetzt analysieren"):
            with st.spinner("Dummy-Analyse läuft…"):
                result = dummy_analyze_cv(
                    content,
                    language,
                    None if selected_role == "Alle" else selected_role,
                    None if selected_industry == "Alle" else selected_industry
                )
            st.success("✅ Analyse abgeschlossen!")
            st.markdown(f"### Ergebnis:\n> {result}")
else:
    st.info("🔔 Bitte lade einen Lebenslauf hoch, um die Analyse zu starten.")
